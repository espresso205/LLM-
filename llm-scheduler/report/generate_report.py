"""
LLM云原生分布式推理引擎 — 项目研究中期报告生成器

格式要求：A4, 页边距2cm, 文档网格46字×43行, 行距16磅,
正文小四号宋体, 图表参考文献五号字, 引注顺序编码制[1][2]上角标。

用法:
    cd llm-scheduler
    python -m report.generate_report
"""
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt

from report.styles import (
    setup_page, add_cover_title, add_cover_info, add_page_break,
    add_author_footer, add_ref_title, add_reference,
)
from report.sections.section1_background import add_section1
from report.sections.section2_plan import add_section2
from report.sections.section3_progress import add_section3
from report.sections.section4_goal import add_section4
from report.sections.section5_funding import add_section5


# ── References (仅文中引注过的) ─────────────────────────────────────────────

REFERENCES = [
    "[1] OpenAI. GPT-4 Technical Report[J]. arXiv preprint arXiv:2303.08774, 2023.",
    "[2] Touvron H, Lavril T, Izacard G, et al. LLaMA: Open and Efficient Foundation "
    "Language Models[J]. arXiv preprint arXiv:2302.13971, 2023.",
    "[3] Kwon W, Li Z, Zhuang S, et al. Efficient Memory Management for Large Language "
    "Model Serving with PagedAttention[C]//SOSP. 2023: 611-626.",
    "[4] Hugging Face. Text Generation Inference[EB/OL]. "
    "https://github.com/huggingface/text-generation-inference, 2024.",
    "[5] NVIDIA. TensorRT-LLM: High-Performance LLM Inference Engine[EB/OL]. "
    "https://github.com/NVIDIA/TensorRT-LLM, 2024.",
    "[6] Newman S. Building Microservices: Designing Fine-Grained Systems[M]. "
    "2nd ed. O'Reilly Media, 2021.",
    "[7] Fielding R T. Architectural Styles and the Design of Network-based "
    "Software Architectures[D]. University of California, Irvine, 2000.",
]


# ── Cover Page ──────────────────────────────────────────────────────────────

def add_cover_page(doc):
    """封面页."""
    add_cover_title(doc, "项目研究中期报告")

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)

    add_cover_title(doc, "LLM云原生分布式推理引擎\n设计与实现")

    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)

    add_cover_info(doc, "学    院", "[学院名称]")
    add_cover_info(doc, "专    业", "[专业名称]")
    add_cover_info(doc, "学生姓名", "[姓名]")
    add_cover_info(doc, "学    号", "[学号]")
    add_cover_info(doc, "指导教师", "[指导教师]")
    add_cover_info(doc, "日    期", "2026年4月")


# ── Main Assembly ───────────────────────────────────────────────────────────

def generate():
    doc = Document()
    setup_page(doc)

    # 封面
    print("[1/7] 封面...")
    add_cover_page(doc)
    add_page_break(doc)

    # 首页地脚 — 作者简介
    add_author_footer(doc, "作者简介：[姓名]，[学院名称][专业名称]专业本科生，研究方向为分布式系统与AI推理。")

    # （一）立项背景
    print("[2/7] （一）立项背景...")
    add_section1(doc)

    # （二）研究内容及实施方案
    print("[3/7] （二）研究内容及实施方案...")
    add_section2(doc)

    # （三）项目进展情况（重点）
    print("[4/7] （三）项目进展情况...")
    add_section3(doc)

    # （四）结题预期目标
    print("[5/7] （四）结题预期目标...")
    add_section4(doc)

    # （五）经费使用情况
    print("[6/7] （五）经费使用情况...")
    add_section5(doc)

    # 参考文献
    print("[7/7] 参考文献...")
    add_page_break(doc)
    add_ref_title(doc, "参考文献")
    for ref in REFERENCES:
        add_reference(doc, ref)

    # 保存
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "mid_term_report.docx")
    doc.save(output_path)
    print(f"\n报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    generate()
