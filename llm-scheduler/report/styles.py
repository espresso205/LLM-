"""Word document style definitions — 中期报告格式。

格式要求：
- A4纸型，页边距各2cm
- 文档网格 46字×43行
- 行距16磅
- 正文小四号宋体，阿拉伯数字/英文 Times New Roman
- 图名、表名及内容、参考文献均为五号字
- 引注采用顺序编码制[1][2]上角标
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Constants ───────────────────────────────────────────────────────────────

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_EN = "Times New Roman"

SIZE_ER = 22       # 二号
SIZE_SAN = 16      # 三号
SIZE_SI = 14       # 四号
SIZE_XIAOSI = 12   # 小四号
SIZE_WU = 10.5     # 五号
LINE_SPACING = Pt(16)


# ── Page Setup ──────────────────────────────────────────────────────────────

def setup_page(doc: Document):
    """A4, 页边距各2cm, 文档网格46字×43行."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 文档网格: 46字×43行
    sectPr = section._sectPr
    docGrid = sectPr.find(qn('w:docGrid'))
    if docGrid is None:
        docGrid = OxmlElement('w:docGrid')
        sectPr.append(docGrid)
    docGrid.set(qn('w:type'), 'linesAndChars')
    docGrid.set(qn('w:linePitch'), '43')
    docGrid.set(qn('w:charSpace'), '0')


# ── Font Helpers ────────────────────────────────────────────────────────────

def _set_run_font(run, cn_font=FONT_SONG, en_font=FONT_EN,
                  size=SIZE_XIAOSI, bold=False, color=None, superscript=False):
    """Set run font with CJK support."""
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = en_font
    # East Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    # Superscript
    if superscript:
        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript')
        rPr.append(vertAlign)


def _set_paragraph_spacing(p, line_spacing=LINE_SPACING,
                           space_before=Pt(0), space_after=Pt(0)):
    """Set paragraph line spacing and spacing."""
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after


# ── Cover Page ──────────────────────────────────────────────────────────────

def add_cover_title(doc, text):
    """封面标题 — 二号黑体加粗居中."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, space_before=Pt(72), space_after=Pt(24))
    run = p.add_run(text)
    _set_run_font(run, cn_font=FONT_HEI, size=SIZE_ER, bold=True)


def add_cover_info(doc, label, value):
    """封面信息行 — 三号宋体居中."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, space_before=Pt(10), space_after=Pt(10))
    run = p.add_run(f"{label}：{value}")
    _set_run_font(run, size=SIZE_SAN)


# ── Author Footer (首页地脚) ───────────────────────────────────────────────

def add_author_footer(doc, text):
    """首页地脚处作者简介 — 五号字."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, space_before=Pt(4), space_after=Pt(4))
    run = p.add_run(text)
    _set_run_font(run, size=SIZE_WU)


# ── Section Headings ────────────────────────────────────────────────────────

def add_section_heading(doc, text):
    """一级标题（一）（二）… — 四号黑体加粗."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, line_spacing=LINE_SPACING,
                           space_before=Pt(12), space_after=Pt(6))
    run = p.add_run(text)
    _set_run_font(run, cn_font=FONT_HEI, size=SIZE_SI, bold=True)


def add_subheading(doc, text):
    """二级标题 1. 2. … — 小四号黑体加粗."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, line_spacing=LINE_SPACING,
                           space_before=Pt(6), space_after=Pt(4))
    run = p.add_run(text)
    _set_run_font(run, cn_font=FONT_HEI, size=SIZE_XIAOSI, bold=True)


# ── Body Text ───────────────────────────────────────────────────────────────

def add_body(doc, text, indent=True):
    """正文 — 小四号宋体, 行距16磅, 首行缩进2字符."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    _set_run_font(run, size=SIZE_XIAOSI)


def add_body_with_cite(doc, parts):
    """带引注的正文段落。

    parts 是一个列表, 每项为:
      - str: 普通文本
      - tuple(str, str): (文本, 'cite') 表示引注上角标如 [1]
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p)
    p.paragraph_format.first_line_indent = Pt(24)
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            _set_run_font(run, size=SIZE_XIAOSI)
        elif isinstance(part, tuple) and part[1] == 'cite':
            run = p.add_run(part[0])
            _set_run_font(run, size=SIZE_XIAOSI, superscript=True)


# ── Table (五号字) ─────────────────────────────────────────────────────────

def add_table(doc, headers, rows, caption=""):
    """表格 — 五号字, 表名居中, 随文出现."""
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, space_before=Pt(8), space_after=Pt(4))
        run = p.add_run(caption)
        _set_run_font(run, cn_font=FONT_HEI, size=SIZE_WU, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(header)
        _set_run_font(run, cn_font=FONT_HEI, size=SIZE_WU, bold=True)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(str(val))
            _set_run_font(run, size=SIZE_WU)
    return table


# ── Figure Placeholder (五号字) ─────────────────────────────────────────────

def add_figure(doc, caption, desc=""):
    """图片占位 — 五号字, 图名居中."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, space_before=Pt(8), space_after=Pt(2))
    run = p.add_run(f"[{desc}]" if desc else "[图片占位]")
    _set_run_font(run, size=SIZE_WU, color=RGBColor(128, 128, 128))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p2, space_before=Pt(2), space_after=Pt(8))
    run2 = p2.add_run(caption)
    _set_run_font(run2, cn_font=FONT_HEI, size=SIZE_WU, bold=True)


# ── References (五号字) ─────────────────────────────────────────────────────

def add_reference(doc, text):
    """参考文献条目 — 五号字."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p, space_before=Pt(1), space_after=Pt(1))
    run = p.add_run(text)
    _set_run_font(run, size=SIZE_WU)


def add_ref_title(doc, text):
    """参考文献标题 — 四号黑体."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, line_spacing=LINE_SPACING,
                           space_before=Pt(12), space_after=Pt(6))
    run = p.add_run(text)
    _set_run_font(run, cn_font=FONT_HEI, size=SIZE_SI, bold=True)


# ── Utility ─────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()
