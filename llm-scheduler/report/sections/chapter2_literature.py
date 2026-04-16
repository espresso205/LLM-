"""第二章 国内外研究现状 —— 文献综述章节。"""

from docx import Document
from report.styles import (
    add_heading1,
    add_heading2,
    add_heading3,
    add_body_text,
    add_table,
    add_page_break,
)


def add_chapter2(doc: Document):
    """生成第二章：国内外研究现状。"""

    # ── 章标题 ──────────────────────────────────────────────────────────────
    add_heading1(doc, "第二章 国内外研究现状")

    # ================================================================== 2.1
    add_heading2(doc, "2.1 LLM推理技术综述")

    add_heading3(doc, "2.1.1 主流推理框架")
    add_body_text(doc,
        "近年来，随着大语言模型参数规模从数十亿跃升至数千亿，推理阶段的计算效率与显存利用率"
        "成为制约其落地的关键瓶颈。围绕这一问题，学术界与工业界相继推出了多种高性能推理框架。"
        "其中具有代表性的包括 vLLM、Text Generation Inference（TGI）、TensorRT-LLM 以及 LMDeploy。"
    )
    add_body_text(doc,
        "vLLM 由加州大学伯克利分校团队提出，其核心创新在于 PagedAttention 机制——借鉴操作系统"
        "虚拟内存分页思想，将 KV Cache 划分为固定大小的 Block 进行非连续存储与按需分配，从而将"
        "显存碎片率从传统方案的 60% 以上降至 4% 以内，显著提升了推理吞吐量。同时，vLLM 引入"
        "Continuous Batching（连续批处理）策略，在每个解码步骤动态调度请求，避免因序列长度差异"
        "导致的计算资源浪费。"
    )
    add_body_text(doc,
        "HuggingFace 推出的 TGI 框架针对生产环境进行了深度优化，支持张量并行（Tensor Parallelism）"
        "与流水线并行（Pipeline Parallelism），并提供连续批处理、Flash Attention 等加速特性。"
        "NVIDIA 的 TensorRT-LLM 则基于 TensorRT 深度学习编译器，针对 NVIDIA GPU 架构进行了极致"
        "的算子级优化，支持 INT8/INT4 量化推理、CUDALLM Kernel 融合以及 Inflight Batching，在"
        "单卡及多卡部署场景下均展现出卓越性能。LMDeploy 由商汤科技开源，集成了推理引擎 TurboMind，"
        "支持 FP16、INT8 及 INT4 多精度推理，在保证精度的前提下大幅降低推理延迟。"
    )

    # ── 表 2-1 LLM推理框架对比 ─────────────────────────────────────────────
    add_table(
        doc,
        headers=["框架", "核心优化技术", "并行策略", "量化支持", "开源协议"],
        rows=[
            ["vLLM", "PagedAttention、Continuous Batching",
             "张量并行", "AWQ/GPTQ", "Apache 2.0"],
            ["TGI", "Flash Attention、Continuous Batching",
             "张量/流水线并行", "bitsandbytes/GPTQ", "Apache 2.0"],
            ["TensorRT-LLM", "Kernel融合、Inflight Batching",
             "张量/流水线并行", "INT8/INT4/FP8", "NVIDIA License"],
            ["LMDeploy", "TurboMind引擎、持久批处理",
             "张量并行", "INT8/INT4", "Apache 2.0"],
        ],
        caption="表2-1 主流LLM推理框架对比",
    )

    add_heading3(doc, "2.1.2 核心优化技术")
    add_body_text(doc,
        "LLM推理优化的关键技术可归纳为以下四个方面：（1）KV Cache 管理——vLLM 的 PagedAttention"
        " 将 KV Cache 分页管理，支持跨请求共享前缀缓存，显著降低显存碎片；（2）连续批处理——"
        "打破传统 Static Batching 的等待壁垒，在每次迭代中动态插入新请求并移除已完成请求，"
        "最大化 GPU 利用率；（3）量化推理——通过 INT8 或 INT4 量化降低模型权重与激活值的存储"
        "和计算开销，结合 GPTQ、AWQ 等训练后量化算法在精度损失可控的前提下实现 2-4 倍吞吐提升；"
        "（4）FlashAttention——通过 IO 感知的注意力计算重排，将中间结果保留在 SRAM 中以减少"
        "HBM 访问次数，在长序列场景下实现接近 2 倍的加速。"
    )

    # ================================================================== 2.2
    add_heading2(doc, "2.2 分布式推理调度研究")

    add_body_text(doc,
        "在分布式推理场景中，如何高效调度多副本模型实例以应对动态变化的请求负载，是推理服务"
        "系统面临的核心挑战。当前业界已有多种成熟的模型服务框架。"
    )
    add_body_text(doc,
        "Ray Serve 是 Ray 分布式计算框架的模型服务组件，采用异构计算资源统一调度架构，支持"
        "自动缩放、流量路由与多模型组合部署，具备良好的通用性，但其主要面向通用 ML 工作负载，"
        "对 LLM 推理特有的 KV Cache 管理与生成式批处理缺乏针对性优化。"
    )
    add_body_text(doc,
        "NVIDIA Triton Inference Server 是业界广泛采用的推理服务方案，支持 TensorFlow、PyTorch、"
        "ONNX 等多种框架模型的统一部署，提供动态批处理与模型版本管理功能，但其架构较为重量级，"
        "配置复杂度较高。KServe 基于 Kubernetes 原生架构构建，支持 Canary 发布、自动缩放等"
        "云原生特性，能够实现模型服务的灰度升级与弹性伸缩，但对 Kubernetes 集群有强依赖。"
        "OpenVINO Model Server 由 Intel 推出，针对 Intel 硬件进行了推理优化，支持模型热更新"
        "与多模型实例调度，但在 GPU 加速场景下的支持相对有限。"
    )

    # ── 表 2-2 分布式推理平台对比 ──────────────────────────────────────────
    add_table(
        doc,
        headers=["平台", "部署方式", "自动缩放", "LLM专项优化", "适用场景"],
        rows=[
            ["Ray Serve", "独立集群", "支持", "无专项优化", "通用ML服务"],
            ["Triton", "Docker/K8s", "支持", "动态批处理", "多框架模型部署"],
            ["KServe", "Kubernetes", "支持", "无专项优化", "云原生模型服务"],
            ["OpenVINO MS", "Docker", "有限", "Intel硬件优化", "CPU推理场景"],
        ],
        caption="表2-2 主流分布式推理平台对比",
    )

    add_body_text(doc,
        "综合对比分析可知，上述方案各有侧重：Ray Serve 通用性强但缺乏 LLM 专项优化；Triton "
        "功能完备但架构重量级；KServe 依赖 Kubernetes 生态；OpenVINO Model Server 聚焦 Intel "
        "硬件。本项目定位为轻量级、专注于 LLM 场景的分布式推理调度方案，旨在弥补现有框架在"
        "易用性与 LLM 场景适配性方面的不足。"
    )

    # ================================================================== 2.3
    add_heading2(doc, "2.3 负载均衡策略研究")

    add_heading3(doc, "2.3.1 经典负载均衡算法")
    add_body_text(doc,
        "负载均衡技术是分布式系统研究的经典课题，其算法演进大致经历三个阶段：第一阶段为静态"
        "算法，以轮询（Round Robin）和加权轮询（Weighted Round Robin）为代表，根据预设权重"
        "分配请求，实现简单但不感知后端实际负载；第二阶段为动态算法，包括最少连接数（Least "
        "Connections）和最短响应时间（Least Response Time）等策略，通过实时监测后端状态做出"
        "调度决策，具有更好的负载适应性；第三阶段为自适应算法，结合机器学习方法对流量模式进行"
        "预测与动态调整，进一步提升资源利用效率。"
    )

    add_heading3(doc, "2.3.2 LLM推理场景的特殊考量")
    add_body_text(doc,
        "LLM推理场景下的负载均衡面临独特挑战：首先，由于自回归生成机制，不同请求的生成长度"
        "差异极大，导致推理延迟呈现长尾分布，传统的静态权重分配难以适应；其次，KV Cache 占用"
        "与当前活跃请求的上下文长度正相关，仅以连接数作为负载指标可能产生偏差；此外，连续批处理"
        "窗口中正在处理的 Token 数量直接影响新请求的排队等待时间。因此，需要综合考虑活跃连接数、"
        "KV Cache 占用量及请求处理延迟等多维指标进行调度决策。"
    )
    add_body_text(doc,
        "本项目实现了三种负载均衡策略：轮询策略适用于后端节点性能均匀的简单场景；最少连接策略"
        "根据各工作节点当前活跃请求数进行动态分配，适用于请求处理时间差异较大的场景；加权轮询"
        "策略允许根据节点硬件配置设置差异化权重，适用于异构集群部署。系统支持策略的运行时热切换，"
        "无需重启服务即可调整调度策略。"
    )

    # ================================================================== 2.4
    add_heading2(doc, "2.4 云原生可观测性")

    add_body_text(doc,
        "云原生环境下的系统可观测性是保障服务可靠性的重要基础，其核心支柱包括指标监控（Metrics）、"
        "分布式追踪（Traces）与日志管理（Logs）三个维度。"
    )
    add_body_text(doc,
        "Prometheus 与 Grafana 组合是当前云原生监控的事实标准。Prometheus 基于 Pull 模式采集"
        "时序指标数据，支持 PromQL 查询语言与灵活的告警规则配置；Grafana 提供丰富的可视化仪表盘"
        "能力，两者结合形成了完善的监控可视化方案。OpenTelemetry 作为 CNCF 的统一可观测性框架，"
        "通过定义标准化的 API 与 SDK 规范，实现了 Traces、Metrics、Logs 三大信号的统一采集与"
        "导出，有效解决了不同厂商可观测性工具之间的数据孤岛问题。"
    )
    add_body_text(doc,
        "然而，完整的 Prometheus+Grafana+OpenTelemetry 技术栈部署与运维成本较高，对于小规模"
        "LLM 推理集群而言存在过度工程化的问题。本项目因此采用自研轻量级监控方案：内置指标采集"
        "模块实时统计推理延迟、吞吐量、GPU 显存使用率等关键指标，提供 Web 可视化界面，同时兼容"
        "Prometheus 指标导出格式，在降低运维复杂度的同时保留了与标准监控生态对接的能力。"
    )

    # ================================================================== 2.5
    add_heading2(doc, "2.5 现有工作不足与本项目定位")

    add_body_text(doc,
        "通过对上述研究领域相关工作的综合分析，现有方案主要存在以下不足：第一，主流推理框架"
        "（vLLM、TGI 等）虽然单引擎推理性能优异，但缺乏开箱即用的分布式多节点调度能力；第二，"
        "通用模型服务平台（Ray Serve、Triton 等）功能完备但架构复杂，学习与部署成本较高，"
        "对中小规模 LLM 推理场景不够友好；第三，现有方案在负载均衡、容错恢复等方面缺乏针对 "
        "LLM 推理特性的精细化优化。"
    )
    add_body_text(doc,
        "基于上述分析，本项目明确自身定位为轻量级、易部署、专注 LLM 推理调度的分布式服务引擎。"
        "其核心创新点包括：（1）策略热切换机制——支持负载均衡策略在运行时动态切换，无需中断"
        "服务即可适应负载模式变化；（2）心跳容错机制——通过定期心跳检测自动剔除故障节点并在"
        "节点恢复后自动注册，保障服务高可用；（3）一体化轻量监控——集成指标采集与 Web 可视化，"
        "降低独立部署监控系统的运维负担。本项目旨在为中小规模 LLM 推理集群提供一套简捷高效的"
        "调度与运维解决方案。"
    )

    add_page_break(doc)
